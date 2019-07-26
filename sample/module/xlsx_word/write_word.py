# -*- coding: utf-8 -*-
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
# font-family
# font-color
# font-size
# underline
# text-align
document = Document()

# font = document.styles['Normal'].font
# font.name = 'SimSun'

def str2unicode(input_str):
    if type(input_str) == str:
        return input_str.decode('utf-8', 'ignore')
    return input_str

def is_number(s):
    try:
        import unicodedata
        unicodedata.numeric(s)
        return True
    except (TypeError, ValueError):
        pass

    return False

def extend_str(input_str, common_len, endfix='\t'):
    result = input_str + ''.join([endfix for x in range((common_len - len(input_str))*2)])
    return result

def get_max_len_4_strs(user_dic=dict()):
    return max([len(str2unicode(user_dic[key])) for key in user_dic])

def set_align(obj, align=0):
    if align == 0:
        obj.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 1:
        obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 2:
        obj.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_underline(run, type=1):
    if type == 0:
        run.font.underline = WD_UNDERLINE.NONE
    elif type == 1:
        run.font.underline = WD_UNDERLINE.SINGLE
    elif type == 2:
        run.font.underline = WD_UNDERLINE.WORDS

def draw_head(content='', level=1, font_family=u'宋体', font_color=(0x3f, 0x2c, 0x36), font_size=26, line_spacing=1, align=0): 
    head = document.add_heading('', level=level)
    head.paragraph_format.line_spacing = line_spacing

    run = head.add_run(content)
    font = run.font
    font.color.rgb = RGBColor(*font_color)
    font.name = font_family
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)
    font.size = Pt(font_size)
    set_align(head, align)

def draw_paragraph(content='', level=1, font_family=u'宋体', font_color=(0x3f, 0x2c, 0x36), font_size=14, underline=0, line_spacing=2, align=0): 
    p = document.add_paragraph('')
    p.paragraph_format.line_spacing = line_spacing

    run = p.add_run(content)
    font = run.font
    font.color.rgb = RGBColor(*font_color)
    font.name = font_family
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)
    font.size = Pt(font_size)
    set_underline(run, underline)
    set_align(p, align)

    return p

def add_paragraph(p, content='', level=1, font_family=u'宋体', font_color=(0x3f, 0x2c, 0x36), font_size=14, underline=0):
    run = p.add_run(content)
    font = run.font
    font.color.rgb = RGBColor(*font_color)
    font.name = font_family
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)
    font.size = Pt(font_size)
    set_underline(run, underline)

def draw_table(table_list=[]):
    table = document.add_table(rows=0, cols=2)
    row = 0
    col = 0
    for row_list in table_list:
        row = table.add_row()
        row.cells[0].text = row_list[0]
        row.cells[1].text = row_list[1]

def draw_index_userinfo(user_dic=dict()):
    # draw_table([
    #     [u'学    院: ', str2unicode(user_dic['college'])],
    #     [u'课程名称: ', str2unicode(user_dic['cozname'])],
    #     [u'专业班级: ', str2unicode(user_dic['class'])],
    #     [u'姓    名: ', str2unicode(user_dic['username'])],
    #     [u'班    级: ', str2unicode(user_dic['number'])]
    # ])
    common_len = get_max_len_4_strs(user_dic)
    # add_paragraph(draw_paragraph(u'学    院：', align=1), extend_str(str2unicode(user_dic['college']), common_len),  underline=1)
    # add_paragraph(draw_paragraph(u'课程名称：', align=1), extend_str(str2unicode(user_dic['cozname']), common_len),  underline=1)
    # add_paragraph(draw_paragraph(u'专业班级：', align=1), extend_str(str2unicode(user_dic['class']), common_len),    underline=1)
    # add_paragraph(draw_paragraph(u'姓    名：', align=1), extend_str(str2unicode(user_dic['username']), common_len), underline=1)
    # add_paragraph(draw_paragraph(u'学    号：', align=1), extend_str(str2unicode(user_dic['number']), common_len),   underline=1)
    add_paragraph(draw_paragraph(u'学    院：', align=1), '\t\t\t\t\t\t', underline=1)
    add_paragraph(draw_paragraph(u'课程名称：', align=1), '\t\t\t\t\t\t', underline=1)
    add_paragraph(draw_paragraph(u'专业班级：', align=1), '\t\t\t\t\t\t', underline=1)
    add_paragraph(draw_paragraph(u'姓    名：', align=1), '\t\t\t\t\t\t', underline=1)
    add_paragraph(draw_paragraph(u'学    号：', align=1), '\t\t\t\t\t\t', underline=1)

draw_head(u'\n\n\n\n江西财经大学',      align=1)
draw_head(u'\n学 生 实 验 报 告\n\n\n', align=1, font_size=36)

userinfo = {
    'college' : '江西财经大学',
    'cozname' : '推断统计学在金融大数据分析中应用的虚拟仿真实验',
    'class'   : '131114',
    'username': '张三',
    'number'  : '13111423'
}

draw_index_userinfo(userinfo)
document.add_page_break()
draw_head(u'一、实验综述 ',                                     level=2, font_size=14)
draw_head(u'1、实验目的及要求 ',                                level=3, font_size=14)
draw_paragraph(u'目的要明确，要抓住重点，符合实验指导书中的要求（使用时斜体字删去，下同）\n', font_size=12)
draw_paragraph(u'===============This is answer===============')
draw_head(u'\n2、实验仪器、设备或软件  ',                       level=3, font_size=14)
draw_paragraph(u'实验所使用的仪器设备、工具等的名称及规格。\n',          font_size=12)
draw_paragraph(u'===============This is answer===============')
draw_head(u'\n二、实验过程（实验步骤、记录、数据、分析）',      level=2, font_size=14)
draw_paragraph(u'写明具体实施的步骤，包括实验过程中的记录、数据和相应的分析\n',               font_size=12)
draw_paragraph(u'===============This is answer===============')
draw_head(u'\n三、结论',                                        level=2, font_size=14)
draw_head(u'1、实验结果',                                       level=3, font_size=14)
draw_paragraph(u'根据实验过程中所见到的现象和测得的数据，做出结论\n',    font_size=12)
draw_paragraph(u'===============This is answer===============')
draw_head(u'\n2、分析讨论',                                     level=3, font_size=14)
draw_paragraph(u'对本次实验的心得体会、思考和建议。\n',                  font_size=12)
draw_paragraph(u'===============This is answer===============')
draw_head(u'\n四、指导教师评语及成绩：',                        level=2, font_size=14)
draw_paragraph(u'评语：指导教师依据学生的实际报告内容，用简练语言给出本次实验报告的评价和价值\n\n\n\n\n\n\n', font_size=12)
draw_paragraph(u'成绩：\t\t指导教师签名：\t\t\t',      align=2)
draw_paragraph(u'批阅日期：\t\t\t',                    align=2)
# p = document.add_paragraph('A plain paragraph having some ')_element.rPr.rFonts
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')

# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )


# records = (
#     (3, '101', 'Spam'),
#     (7, '422', 'Eggs'),
#     (4, '631', 'Spam, spam, eggs, and spam')
# )

# table = document.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# for qty, id, desc in records:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(qty)
#     row_cells[1].text = id
#     row_cells[2].text = desc

# document.add_page_break()

document.save('demo.docx')